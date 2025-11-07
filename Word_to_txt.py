from docx import Document


def extract_text_from_docx(file_path):
    # 加载文档
    doc = Document(file_path)

    # 存储提取的文本
    full_text = []

    # 1. 提取段落文本
    for para in doc.paragraphs:
        # 段落文本可能为空（如空行），过滤掉纯空格
        if para.text.strip():
            full_text.append(para.text)

    # 2. 提取表格文本（遍历所有表格、行、单元格）
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                # 单元格文本可能为空，过滤纯空格
                if cell.text.strip():
                    full_text.append(cell.text)

    # 将所有文本用换行符连接（保持原文档的大致结构）
    return '\n'.join(full_text)


# 使用示例
if __name__ == '__main__':
    docx_path = "F:\第一部分（丄至竟）.docx"  # 替换为你的.docx文件路径
    text = extract_text_from_docx(docx_path)
    print(text)  # 打印提取的文本
    # 也可以保存到txt文件
    with open('F:\\1.txt', 'w', encoding='utf-8') as f:
        f.write(text)